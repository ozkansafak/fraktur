# document-generation/utils.py

import requests
import logging
import os
from io import BytesIO

def upload_document(file_path: str, config) -> str:
    """
    Uploads the generated document to the Storage Service and returns the file URL.
    
    Args:
        file_path (str): Path to the generated .docx file.
        config (Config): Configuration object containing STORAGE_SERVICE_URL.
    
    Returns:
        str: URL of the uploaded document.
    """
    try:
        with open(file_path, 'rb') as f:
            files = {'file': (os.path.basename(file_path), f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            response = requests.post(config.STORAGE_SERVICE_URL, files=files)
        response.raise_for_status()
        file_url = response.json().get('file_url')
        return file_url
    except requests.exceptions.RequestException as e:
        logging.error(f"Failed to upload document to Storage Service: {e}")
        raise

def create_document(german_text: str, english_text: str, config, output_path: str) -> None:
    """
    Creates a .docx document with German and English texts.
    
    Args:
        german_text (str): Extracted German transcription.
        english_text (str): Translated English text.
        config (Config): Configuration object for styling and storage.
        output_path (str): Path to save the generated .docx file.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    document = Document()

    # Add German Transcription
    document.add_heading('German Transcription', level=1)
    p = document.add_paragraph(german_text)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)

    # Add a page break
    document.add_page_break()

    # Add English Translation
    document.add_heading('English Translation', level=1)
    p = document.add_paragraph(english_text)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)

    # Save the document
    document.save(output_path)
