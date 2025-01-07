from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
import re
import os
import PyPDF2
import logging
from dataclasses import dataclass
from typing import Tuple, Dict, Any, Optional
from src.utils import setup_logger

logger = logging.getLogger('logger_name')
logger.setLevel(logging.INFO)


def chapter_splitter(input_pdf_path, output_folder):
    logger = setup_logger('time_logger')
    
    # Ensure the output folder exists
    os.makedirs(output_folder, exist_ok=True)
    
    # Open the PDF and split pages
    try:
        with open(input_pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            logger.info(f"Extracted {num_pages} pages from PDF")
    
            for i, page in enumerate(pdf_reader.pages):
                # Create a new PDF writer for each page
                pdf_writer = PyPDF2.PdfWriter()
                pdf_writer.add_page(page)
    
                # Save the current page to a new file
                output_file_path = os.path.join(output_folder, f"page_{i+1:03d}.pdf")
                with open(output_file_path, 'wb') as output_file:
                    pdf_writer.write(output_file)
                
                logger.debug(f"Saved page {i+1} to {output_file_path}")
    
        logger.info(f"All pages have been split and saved to {output_folder}")
    
    except Exception as e:
        logger.error(f"An error occurred: {e}")


def strip_newlines(text: str) -> str:
    """Clean text by removing newlines adjacent to section tags."""
    # Remove \n before and after section tags
    text = re.sub(r'\n*<(body|header|footer)>', r'<\1>', text)
    text = re.sub(r'</(body|header|footer)>\n*', r'</\1>', text)
    
    return text

def extract_sections_in_order(text: str) -> list:
    """
    Extracts sections in the order they appear in the text.
    
    Args:
        text (str): Text containing header, body, and footer tags
        
    Returns:
        list: List of tuples (section_type, section_content) in order of appearance
    """
    # Pattern to match any section
    pattern = r'<(pageno|header|body|footer)>(.*?)</\1>'
    
    # Find all sections in order of appearance
    sections = []
    for match in re.finditer(pattern, text, re.DOTALL):
        section_type = match.group(1)
        content = match.group(2)
        sections.append((section_type, content))
    
    return sections


def add_tab_stop(paragraph, position_in_inches):
    """Adds a right-aligned tab stop to the paragraph."""
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(position_in_inches), alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT)

def add_bottom_border(paragraph: Any) -> None:
    """Adds a bottom border to a paragraph."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    borders = pPr.find(qn('w:pBdr'))
    if borders is None:
        borders = OxmlElement('w:pBdr')
        pPr.append(borders)
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '2')
    bottom_border.set(qn('w:space'), '1')
    bottom_border.set(qn('w:color'), 'auto')
    borders.append(bottom_border)

def setup_document_styles(document: Document) -> Dict[str, Any]:
    """Sets up document styles and returns them."""
    styles = document.styles

    # Header style
    if 'CustomHeaderStyle' not in styles:
        header_style = styles.add_style('CustomHeaderStyle', WD_STYLE_TYPE.PARAGRAPH)
        header_style.font.name = 'Arial'
        header_style.font.size = Pt(14)
        header_style.font.color.rgb = RGBColor(0, 0, 0)

    # Footnote style
    if 'FootnoteStyle' not in styles:
        footnote_style = styles.add_style('FootnoteStyle', WD_STYLE_TYPE.PARAGRAPH)
        footnote_style.font.size = Pt(9)

    return {
        'header': styles['CustomHeaderStyle'],
        'footnote': styles['FootnoteStyle']
    }

def save_document(texts: dict, 
                  folder_name: str = '', 
                  language: str = 'English') -> Tuple[Document, str]:
    """
    Creates a .docx document maintaining the original section order.
    """
    logger = setup_logger('ocr_processor')

    document = Document()
    styles = setup_document_styles(document)

    # Configure document margins
    section = document.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Process each page
    for index, pageno in enumerate(sorted(texts.keys())):
        text = re.sub(r'\n+', '\n', texts[pageno])
        
        text = strip_newlines(text)
        
        # Add page break for all pages except the first
        if index > 0:
            document.add_page_break()

        # Get sections in their original order
        sections = extract_sections_in_order(text)

        # Check the first section type
        if sections and sections[0][0] == 'pageno':
            # Handle case where <pageno> is the first section type
            content = sections.pop(0)[1]  # Remove <pageno> and get content
            header_para = document.add_paragraph(f"Page: {content}\tkeyno: {pageno}")
        else:
            # Handle case where <pageno> is not the first section type
            header_para = document.add_paragraph(f"\tkeyno: {pageno}")
        
        border_para = document.add_paragraph()
        add_bottom_border(border_para)

        # Apply style and alignment for the first line
        header_para.style = styles['header']
        tab_stop_position = section.page_width.inches - section.right_margin.inches
        add_tab_stop(header_para, tab_stop_position)

        for i, (section_type, content) in enumerate(sections):

            if section_type == 'header':
                # Add header content
                header_para = document.add_paragraph(content)
                header_para.style = styles['header']
                header_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            elif section_type == 'body':
                # Add body paragraphs
                for para_text in content.split('\n'):
                    if para_text.strip():
                        body_para = document.add_paragraph(para_text)
                        body_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            elif section_type == 'footer':
                # Add border paragraph before footer
                border_para = document.add_paragraph()
                add_bottom_border(border_para)

                # Add footer content
                footer_para = document.add_paragraph(content)
                footer_para.style = styles['footnote']
                footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Save the document
    fname = f'../output_data/{folder_name}/{language}'
    document.save(f'{fname}.docx')

    return document, fname
